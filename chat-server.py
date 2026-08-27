import http.server
import socketserver
import os
import json
import io

PORT = 3001
MAX_UPLOAD = 50 * 1024 * 1024      # 50 MB body cap
MAX_EXTRACT_CHARS = 100_000        # extraction text cap (protects LLM context)

EXTRACTABLE = {".pdf", ".xlsx", ".xls", ".docx", ".txt", ".csv", ".json", ".md"}


def extract_text(filename, data):
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    note = ""
    if ext == ".pdf":
        from pypdf import PdfReader
        r = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() or "" for p in r.pages]
        text = "\n\n".join(f"[page {i + 1}]\n{t}" for i, t in enumerate(pages))
        note = f"{len(r.pages)} pages"
    elif ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join("" if v is None else str(v) for v in row))
            parts.append(f"[sheet: {ws.title}]\n" + "\n".join(rows))
        text = "\n\n".join(parts)
        note = f"{len(wb.worksheets)} sheets"
        wb.close()
    elif ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        parts = []
        for sh in wb.sheets():
            rows = []
            for r in range(sh.nrows):
                rows.append("\t".join(str(sh.cell_value(r, c)) for c in range(sh.ncols)))
            parts.append(f"[sheet: {sh.name}]\n" + "\n".join(rows))
        text = "\n\n".join(parts)
        note = f"{wb.nsheets} sheets"
    elif ext == ".docx":
        from docx import Document
        d = Document(io.BytesIO(data))
        paras = [p.text for p in d.paragraphs]
        tables = []
        for t in d.tables:
            for row in t.rows:
                tables.append("\t".join(c.text for c in row.cells))
        text = "\n".join(paras + tables)
        note = f"{len(d.paragraphs)} paragraphs"
    else:  # txt / csv / json / md — client usually reads these, but support anyway
        text = data.decode("utf-8", errors="replace")
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + "\n…[truncated]"
    return text, note


class Handler(http.server.SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=os.path.dirname(os.path.abspath(__file__)), **kwargs)

    def end_headers(self):
        # no-store: the browser must NEVER keep index.html (all CSS/JS is
        # inline), so every plain reload re-fetches — no Ctrl+Shift+R, and
        # David never has to clear browsing data (which wiped his chats once).
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        super().end_headers()

    def log_message(self, format, *args):
        # keep the access log ON (2026-08-27): needed to prove whether a
        # browser actually fetched the page (stale-tab debugging) — the old
        # `pass` made "why is my profile showing an old page?" undiagnosable.
        # Guard the write: if the parent session died and stdout is an
        # orphaned pipe, an uncaught BrokenPipeError here kills the request
        # handler (empty reply on every request — hit live 2026-08-27).
        try:
            print(f"[{self.log_date_time_string()}] {self.client_address[0]} {format % args}", flush=True)
        except (BrokenPipeError, OSError):
            pass

    def do_POST(self):
        if self.path != "/api/extract":
            self._json({"error": "not found"}, 404)
            return
        length = int(self.headers.get("Content-Length", 0) or 0)
        if length <= 0 or length > MAX_UPLOAD:
            self._json({"error": f"body must be 1..{MAX_UPLOAD} bytes"},
                       413 if length > MAX_UPLOAD else 400)
            return
        filename = self.headers.get("X-Filename", "file")
        data = self.rfile.read(length)
        ext = os.path.splitext(filename)[1].lower()
        if ext not in EXTRACTABLE:
            self._json({"error": f"unsupported type: {ext or '(none)'}"}, 415)
            return
        try:
            text, note = extract_text(filename, data)
            self._json({"text": text, "filename": filename, "note": note,
                        "chars": len(text),
                        "truncated": len(text) >= MAX_EXTRACT_CHARS})
        except Exception as e:
            self._json({"error": f"extract failed: {e}"}, 422)

    def _json(self, obj, code=200):
        body = json.dumps(obj).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


with http.server.ThreadingHTTPServer(("", PORT), Handler) as httpd:
    print(f"Chat server running on port {PORT}")
    httpd.serve_forever()
